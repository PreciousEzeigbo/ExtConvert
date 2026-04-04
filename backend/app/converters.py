"""
ConversionManager — routes (source_ext, target_format) to the right handler.

Supported routes
────────────────
Image  → PDF  : PIL → reportlab
Image  → DOCX : PIL embed into python-docx
Image  → PNG/JPG/WebP : PIL re-save
Image  → TXT  : pytesseract OCR
PDF    → PNG/JPG/WebP : pdf2image
PDF    → TXT  : pypdf text extraction
PDF    → DOCX : extract text → write docx
DOCX   → TXT  : python-docx paragraph extraction
DOCX   → PNG  : docx→pdf (libreoffice if available) → pdf2image, fallback text render
TXT    → PDF  : reportlab
TXT    → DOCX : python-docx
"""

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

# ── Image handling ─────────────────────────────────────────────────────────────
from PIL import Image

# ── PDF ────────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfgen import canvas

import pypdf
from pdf2image import convert_from_path

# ── DOCX ───────────────────────────────────────────────────────────────────────
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── OCR ────────────────────────────────────────────────────────────────────────
try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".gif"}


class ConversionManager:
    def __init__(self, upload_dir: Path, output_dir: Path):
        self.upload_dir = upload_dir
        self.output_dir = output_dir

    # ──────────────────────────────────────────────────────────────────────────
    # Public entry point
    # ──────────────────────────────────────────────────────────────────────────
    def convert(self, src: str, src_ext: str, target: str, file_id: str) -> str:
        """
        Convert *src* file to *target* format.
        Returns the absolute path of the output file.
        """
        src_ext = src_ext.lower().lstrip(".")
        target  = target.lower().lstrip(".")

        out_path = str(self.output_dir / f"{file_id}.{target}")

        # ── Image → * ─────────────────────────────────────────────────────────
        if src_ext in {e.lstrip(".") for e in IMAGE_EXTS}:
            if target == "pdf":
                return self._image_to_pdf(src, out_path)
            elif target == "docx":
                return self._image_to_docx(src, out_path)
            elif target == "txt":
                return self._image_to_txt(src, out_path)
            elif target in {"png", "jpg", "jpeg", "webp"}:
                return self._image_to_image(src, out_path, target)

        # ── PDF → * ───────────────────────────────────────────────────────────
        elif src_ext == "pdf":
            if target in {"png", "jpg", "jpeg", "webp"}:
                return self._pdf_to_image(src, out_path, target)
            elif target == "txt":
                return self._pdf_to_txt(src, out_path)
            elif target == "docx":
                return self._pdf_to_docx(src, out_path)

        # ── DOCX → * ──────────────────────────────────────────────────────────
        elif src_ext == "docx":
            if target == "txt":
                return self._docx_to_txt(src, out_path)
            elif target in {"png", "jpg", "jpeg"}:
                return self._docx_to_image(src, out_path, target)
            elif target == "pdf":
                return self._docx_to_pdf(src, out_path)

        # ── TXT → * ───────────────────────────────────────────────────────────
        elif src_ext == "txt":
            if target == "pdf":
                return self._txt_to_pdf(src, out_path)
            elif target == "docx":
                return self._txt_to_docx(src, out_path)

        raise ValueError(f"Unsupported conversion: .{src_ext} → .{target}")

    # ──────────────────────────────────────────────────────────────────────────
    # Image converters
    # ──────────────────────────────────────────────────────────────────────────
    def _image_to_pdf(self, src: str, out: str) -> str:
        img = Image.open(src).convert("RGB")
        w, h = img.size

        # reportlab canvas approach – preserves aspect ratio on A4
        a4_w, a4_h = A4
        margin = 1 * cm
        max_w = a4_w - 2 * margin
        max_h = a4_h - 2 * margin
        ratio = min(max_w / w, max_h / h)
        draw_w, draw_h = w * ratio, h * ratio
        x = (a4_w - draw_w) / 2
        y = (a4_h - draw_h) / 2

        c = canvas.Canvas(out, pagesize=A4)
        c.drawImage(src, x, y, width=draw_w, height=draw_h, preserveAspectRatio=True)
        c.save()
        return out

    def _image_to_docx(self, src: str, out: str) -> str:
        doc = Document()
        doc.add_heading("Converted Image", level=1)
        doc.add_picture(src, width=Inches(5))
        doc.save(out)
        return out

    def _image_to_txt(self, src: str, out: str) -> str:
        if not OCR_AVAILABLE:
            raise RuntimeError("pytesseract not installed; OCR unavailable.")
        img = Image.open(src)
        text = pytesseract.image_to_string(img)
        Path(out).write_text(text, encoding="utf-8")
        return out

    def _image_to_image(self, src: str, out: str, fmt: str) -> str:
        img = Image.open(src)
        if fmt in {"jpg", "jpeg"}:
            img = img.convert("RGB")
            img.save(out, "JPEG", quality=92)
        elif fmt == "png":
            img.save(out, "PNG")
        elif fmt == "webp":
            img.save(out, "WEBP", quality=90)
        return out

    # ──────────────────────────────────────────────────────────────────────────
    # PDF converters
    # ──────────────────────────────────────────────────────────────────────────
    def _pdf_to_image(self, src: str, out: str, fmt: str) -> str:
        pages = convert_from_path(src, dpi=150)
        if not pages:
            raise ValueError("PDF has no renderable pages.")

        if len(pages) == 1:
            img = pages[0]
            save_fmt = "JPEG" if fmt in {"jpg", "jpeg"} else ("WEBP" if fmt == "webp" else "PNG")
            img.save(out, save_fmt)
        else:
            # Multi-page: stitch vertically
            widths  = [p.width  for p in pages]
            heights = [p.height for p in pages]
            total_h = sum(heights)
            max_w   = max(widths)
            combined = Image.new("RGB", (max_w, total_h), (255, 255, 255))
            y_off = 0
            for pg in pages:
                combined.paste(pg, (0, y_off))
                y_off += pg.height
            save_fmt = "JPEG" if fmt in {"jpg", "jpeg"} else ("WEBP" if fmt == "webp" else "PNG")
            combined.save(out, save_fmt)
        return out

    def _pdf_to_txt(self, src: str, out: str) -> str:
        reader = pypdf.PdfReader(src)
        lines = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            lines.append(f"── Page {i + 1} ──\n{text}\n")
        Path(out).write_text("\n".join(lines), encoding="utf-8")
        return out

    def _pdf_to_docx(self, src: str, out: str) -> str:
        reader = pypdf.PdfReader(src)
        doc = Document()
        doc.add_heading("Converted PDF", level=1)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            doc.add_heading(f"Page {i + 1}", level=2)
            for para in text.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(out)
        return out

    # ──────────────────────────────────────────────────────────────────────────
    # DOCX converters
    # ──────────────────────────────────────────────────────────────────────────
    def _docx_to_txt(self, src: str, out: str) -> str:
        doc = Document(src)
        lines = [para.text for para in doc.paragraphs]
        Path(out).write_text("\n".join(lines), encoding="utf-8")
        return out

    def _docx_to_pdf(self, src: str, out: str) -> str:
        # Try LibreOffice headless first
        lo = shutil.which("libreoffice") or shutil.which("soffice")
        if lo:
            subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", "--outdir",
                 str(self.output_dir), src],
                check=True, capture_output=True,
            )
            # LibreOffice names it after the source file
            gen = self.output_dir / (Path(src).stem + ".pdf")
            if gen.exists() and str(gen) != out:
                gen.rename(out)
            return out

        # Fallback: extract text → PDF
        doc = Document(src)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        tmp_txt = src + ".tmp.txt"
        Path(tmp_txt).write_text("\n".join(texts), encoding="utf-8")
        result = self._txt_to_pdf(tmp_txt, out)
        Path(tmp_txt).unlink(missing_ok=True)
        return result

    def _docx_to_image(self, src: str, out: str, fmt: str) -> str:
        # Convert docx→pdf first, then pdf→image
        tmp_pdf = src + ".tmp.pdf"
        self._docx_to_pdf(src, tmp_pdf)
        result = self._pdf_to_image(tmp_pdf, out, fmt)
        Path(tmp_pdf).unlink(missing_ok=True)
        return result

    # ──────────────────────────────────────────────────────────────────────────
    # TXT converters
    # ──────────────────────────────────────────────────────────────────────────
    def _txt_to_pdf(self, src: str, out: str) -> str:
        text = Path(src).read_text(encoding="utf-8", errors="replace")
        doc = SimpleDocTemplate(
            out,
            pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2 * cm,  bottomMargin=2 * cm,
        )
        styles = getSampleStyleSheet()
        style = ParagraphStyle(
            "body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
        )
        story = []
        for line in text.split("\n"):
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line or "&nbsp;", style))
            story.append(Spacer(1, 2))
        doc.build(story)
        return out

    def _txt_to_docx(self, src: str, out: str) -> str:
        text = Path(src).read_text(encoding="utf-8", errors="replace")
        doc = Document()
        doc.add_heading("Converted Document", level=1)
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(out)
        return out